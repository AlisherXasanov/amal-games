#!/usr/bin/env python3
"""Generate printable Word sheet: 6 meme QR cards per A4 page."""

from __future__ import annotations

import io
from pathlib import Path

import qrcode
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

SITE_URL = "https://alisherxasanov.github.io/amal-games/go-memes.html?v=1"
SOURCE_PAGE = "https://alisherxasanov.github.io/amal-games/qr-memes.html?print=1"
CARDS_PER_PAGE = 6
ROWS, COLS = 3, 2


def make_qr_png_bytes(url: str, box_size: int = 8) -> bytes:
    qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_M, box_size=box_size, border=2)
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), kwargs[edge].get("val", "single"))
            tag.set(qn("w:sz"), str(kwargs[edge].get("sz", 8)))
            tag.set(qn("w:space"), "0")
            tag.set(qn("w:color"), kwargs[edge].get("color", "229ED9"))
            tc_borders.append(tag)
    tc_pr.append(tc_borders)


def add_paragraph(cell, text: str, *, size: float = 9, bold: bool = False, color: RGBColor | None = None, space_after: float = 2) -> None:
    p = cell.paragraphs[0] if cell.paragraphs and cell.paragraphs[0].text == "" and len(cell.paragraphs) == 1 else cell.add_paragraph()
    if cell.paragraphs[0] is not p and cell.paragraphs[0].text == "" and len(cell.paragraphs) > 1:
        pass
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)


def fill_card(cell, qr_bytes: bytes) -> None:
    cell.text = ""
    border = {"val": "single", "sz": 10, "color": "229ED9"}
    set_cell_border(cell, top=border, left=border, bottom=border, right=border)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Бесплатно · без скачивания")
    r.font.size = Pt(7.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("😂 Мемы Амаля")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x22, 0x9E, 0xD9)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Что случилось дома · смешняшки · новости")
    r.font.size = Pt(7.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(io.BytesIO(qr_bytes), width=Mm(28))

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("📷 Наведи камеру сюда")
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x22, 0x9E, 0xD9)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SITE_URL)
    r.font.size = Pt(5.5)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    steps = [
        "1. Открой камеру на телефоне",
        "2. Наведи на QR — нажми ссылку",
        "3. Напиши имя и кидай мемы",
    ]
    for step in steps:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(step)
        r.font.size = Pt(6.5)
        r.font.bold = True

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Один канал для всех")
    r.font.size = Pt(6.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— Амаль ♥")
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xE2, 0x5A, 0x3C)


def build_document(output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = intro.add_run("QR · Мемы Амаля — лист для печати (6 наклеек на странице)")
    run.bold = True
    run.font.size = Pt(11)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Источник: {SOURCE_PAGE}")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    qr_bytes = make_qr_png_bytes(SITE_URL)
    table = doc.add_table(rows=ROWS, cols=COLS)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    usable_width = section.page_width - section.left_margin - section.right_margin
    col_width = int(usable_width / COLS)
    for col in table.columns:
        col.width = col_width

    for row in table.rows:
        row.height = Cm(8.8)
        for cell in row.cells:
            fill_card(cell, qr_bytes)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("💻 На компьютере: Игры Амаля → Сканер QR")
    r.font.size = Pt(8)
    r.font.bold = True

    doc.save(output_path)


if __name__ == "__main__":
    out = Path("/workspace/docs/qr-memes-6-per-page.docx")
    build_document(out)
    artifact = Path("/opt/cursor/artifacts/qr-memes-6-per-page.docx")
    artifact.write_bytes(out.read_bytes())
    print(f"Saved: {out}")
    print(f"Artifact: {artifact}")
